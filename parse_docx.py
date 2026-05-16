# -*- coding: utf-8 -*-
"""Parse Word tickets into JSON."""
import os
import re
import json
from docx import Document

def get_paragraph_text_and_bold(p):
    """Return full text and whether paragraph has any bold content (excluding question number line logic later)."""
    parts = []
    bold_parts = []
    for r in p.runs:
        t = r.text
        if not t:
            continue
        parts.append(t)
        if r.bold:
            bold_parts.append(t)
    full = ''.join(parts)
    has_bold = len(bold_parts) > 0
    bold_text = ''.join(bold_parts)
    return full.strip(), has_bold, bold_text.strip()

def is_question_line(text):
    return bool(re.match(r'^\d+\.\s', text))

def is_answer_line(text):
    return bool(re.match(r'^[а-гa-dА-ГA-D]\)\s*', text, re.I))

def get_answer_letter(text):
    m = re.match(r'^([а-гa-dА-ГA-D])\)', text, re.I)
    return m.group(1).lower() if m else None

def parse_docx(path):
    doc = Document(path)
    questions = []
    current_q = None
    current_answers = []
    current_answer = None

    for p in doc.paragraphs:
        text, has_bold, bold_text = get_paragraph_text_and_bold(p)
        if not text:
            continue

        # Skip header lines (ticket title without question number at start after digits)
        if re.match(r'^Билет\s', text, re.I) or re.match(r'^Вопросы\s', text, re.I):
            continue

        if is_question_line(text):
            # Save previous question
            if current_q is not None:
                if current_answer is not None:
                    current_answers.append(current_answer)
                    current_answer = None
                questions.append({
                    'question': current_q['question'],
                    'answers': current_answers
                })
            # New question
            qnum_match = re.match(r'^(\d+)\.\s*(.*)$', text, re.DOTALL)
            qnum = int(qnum_match.group(1))
            qtext = qnum_match.group(2).strip()
            current_q = {'num': qnum, 'question': qtext}
            current_answers = []
            current_answer = None
            continue

        if current_q is None:
            continue

        if is_answer_line(text):
            if current_answer is not None:
                current_answers.append(current_answer)
            letter = get_answer_letter(text)
            ans_text = re.sub(r'^[а-гa-dА-ГA-D]\)\s*', '', text, flags=re.I).strip()
            # Determine if this answer is correct: bold in answer line
            # If entire line has bold, or bold portion covers answer content
            correct = has_bold
            current_answer = {
                'letter': letter,
                'text': ans_text,
                'correct': correct
            }
        else:
            # Continuation line - append to question or last answer
            if current_answer is not None:
                current_answer['text'] += ' ' + text
                if has_bold:
                    current_answer['correct'] = True
            else:
                current_q['question'] += ' ' + text
                if has_bold and not is_question_line(text):
                    pass  # question continuation rarely bold for answers

    if current_q is not None:
        if current_answer is not None:
            current_answers.append(current_answer)
        questions.append({
            'question': current_q['question'],
            'answers': current_answers
        })

    return questions

def sort_key_filename(fname):
    m = re.search(r'(\d+)', fname)
    return int(m.group(1)) if m else 0

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    files = [f for f in os.listdir(base) if f.endswith('.docx')]
    # Sort by first number in filename
    files.sort(key=lambda f: (
        int(re.search(r'(\d+)', f).group(1)) if re.search(r'(\d+)', f) else 0
    ))

    all_questions = []
    seen_nums = set()

    for fname in files:
        path = os.path.join(base, fname)
        parsed = parse_docx(path)
        print(f'{fname}: {len(parsed)} questions', file=__import__('sys').stderr)
        for q in parsed:
            all_questions.append(q)

    # Assign IDs and tickets based on question order
    # We need global numbering - extract from question text in files or sequential
    # Re-parse with question numbers
    result = []
    id_counter = 1

    # Re-parse all with numbers
    all_with_nums = []
    for fname in files:
        path = os.path.join(base, fname)
        doc = Document(path)
        current_q = None
        current_answers = []
        current_answer = None

        for p in doc.paragraphs:
            text, has_bold, _ = get_paragraph_text_and_bold(p)
            if not text:
                continue
            if re.match(r'^Билет\s', text, re.I):
                continue

            if is_question_line(text):
                if current_q is not None:
                    if current_answer is not None:
                        current_answers.append(current_answer)
                    all_with_nums.append({
                        'num': current_q['num'],
                        'question': current_q['question'],
                        'answers': current_answers
                    })
                qnum_match = re.match(r'^(\d+)\.\s*(.*)$', text, re.DOTALL)
                current_q = {'num': int(qnum_match.group(1)), 'question': qnum_match.group(2).strip()}
                current_answers = []
                current_answer = None
            elif current_q is not None:
                if is_answer_line(text):
                    if current_answer is not None:
                        current_answers.append(current_answer)
                    letter = get_answer_letter(text)
                    ans_text = re.sub(r'^[а-гa-dА-ГA-D]\)\s*', '', text, flags=re.I).strip()
                    current_answer = {'letter': letter, 'text': ans_text, 'correct': has_bold}
                else:
                    if current_answer is not None:
                        current_answer['text'] += ' ' + text
                        if has_bold:
                            current_answer['correct'] = True
                    else:
                        current_q['question'] += ' ' + text

        if current_q is not None:
            if current_answer is not None:
                current_answers.append(current_answer)
            all_with_nums.append({
                'num': current_q['num'],
                'question': current_q['question'],
                'answers': current_answers
            })

    all_with_nums.sort(key=lambda x: x['num'])

    for q in all_with_nums:
        ticket = (q['num'] - 1) // 20 + 1
        answers = []
        for a in q['answers']:
            answers.append({
                'text': a['text'],
                'correct': a['correct']
            })
        result.append({
            'id': q['num'],
            'ticket': ticket,
            'question': q['question'],
            'answers': answers
        })

    out_path = os.path.join(base, 'questions.json')
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f'Total: {len(result)} questions')
    print(f'Tickets: {max(x["ticket"] for x in result)}')
    # Stats
    multi = sum(1 for q in result if sum(1 for a in q['answers'] if a['correct']) > 1)
    print(f'Multi-correct: {multi}')
    no_correct = [q for q in result if sum(1 for a in q['answers'] if a['correct']) == 0]
    print(f'No correct marked: {len(no_correct)}')
    if no_correct[:3]:
        for q in no_correct[:3]:
            print(f'  Q{q["id"]}: {q["question"][:60]}...')

if __name__ == '__main__':
    main()
